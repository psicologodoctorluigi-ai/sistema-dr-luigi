import streamlit as st
import pandas as pd
from datetime import datetime, timedelta, timezone
import random
import string
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from streamlit_gsheets import GSheetsConnection
import plotly.express as px

# -------------------------
# CONFIGURACIÓN Y ZONA HORARIA
# -------------------------
st.set_page_config(page_title="Consultorio del Dr. Luigi's", layout="wide")

# Forzar zona horaria de Lima, Perú (UTC -5)
ZONA_PERU = timezone(timedelta(hours=-5))

# Nombres de las columnas exactas que tendrá tu Google Sheet
COLUMNAS = [
    "Código", "Fecha", "Hora", "Nombres y Apellidos", "DNI", "Edad", "Sexo", "Cargo", 
    "Área", "Tiempo de servicio", "Tipo de contrato", "Teléfono", "Motivo de Atención", "Solicitante", 
    "Descripción", "Tiempo del problema", "Ámbito del problema", "Actitud", "Observaciones conductuales", 
    "Área afectada", "Orientación", "Acuerdos", "Plan de Acción", "Requiere cita", "Fecha próxima cita", "Conclusión"
]

# -------------------------
# CONEXIÓN A GOOGLE SHEETS Y LIMPIEZA
# -------------------------
def cargar_datos():
    try:
        conn = st.connection("gsheets", type=GSheetsConnection)
        df = conn.read(worksheet="Hoja 1", usecols=list(range(len(COLUMNAS))), ttl=0)
        df = df.dropna(how="all") 
        
        if "DNI" in df.columns:
            df["DNI"] = df["DNI"].astype(str)
            df["DNI"] = df["DNI"].str.replace(r"\.0$", "", regex=True)
            df["DNI"] = df["DNI"].str.replace("'", "", regex=False).str.strip()
            df["DNI"] = df["DNI"].apply(lambda x: x.zfill(8) if x.isdigit() else x)
            
        return df
    except Exception:
        return pd.DataFrame(columns=COLUMNAS)

def guardar_datos(nuevos_datos_dict):
    df_actual = cargar_datos()
    df_nuevo = pd.DataFrame([nuevos_datos_dict])
    df_actualizado = pd.concat([df_actual, df_nuevo], ignore_index=True)
    
    try:
        conn = st.connection("gsheets", type=GSheetsConnection)
        conn.update(worksheet="Hoja 1", data=df_actualizado)
        return True
    except Exception as e:
        st.error(f"⚠️ Aún falta configurar la conexión a Google Sheets.")
        return False

# -------------------------
# FUNCIONES AUXILIARES
# -------------------------
def generar_codigo():
    fecha = datetime.now(ZONA_PERU).strftime("%Y%m%d")
    rand = ''.join(random.choices(string.ascii_uppercase + string.digits, k=4))
    return f"HC-{fecha}-{rand}"

def obtener_codigo_por_dni(dni, df):
    dni_limpio = str(dni).strip().zfill(8)
    if not df.empty and dni_limpio in df["DNI"].values:
        codigo = df[df["DNI"] == dni_limpio]["Código"].iloc[0]
        return codigo, True
    return generar_codigo(), False

def generar_word_memoria(datos):
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.color.rgb = RGBColor(0, 0, 0)
    
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run("FICHA DE ATENCIÓN DE CONSEJERÍA PSICOLÓGICA LABORAL")
    run_titulo.bold = True
    run_titulo.font.size = Pt(14)
    
    sub_titulo = doc.add_paragraph()
    sub_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_titulo.paragraph_format.space_after = Pt(20)
    run_sub = sub_titulo.add_run("Subgerencia de Seguridad Ciudadana")
    run_sub.bold = True
    run_sub.font.size = Pt(12)

    def add_subtitulo(texto, numero):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(14) 
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{numero}. {texto}")
        run.bold = True
        run.font.size = Pt(12)

    def add_texto(etiqueta, valor):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6) 
        
        run_etiq = p.add_run(f"{etiqueta}: ")
        run_etiq.bold = True
        run_etiq.font.size = Pt(11)
        
        if etiqueta == "Edad":
            try:
                val_str = str(int(float(valor)))
            except:
                val_str = str(valor)
        elif etiqueta == "DNI":
            val_str = str(valor).replace("'", "")
        else:
            val_str = str(valor)
            
        run_val = p.add_run(val_str)
        run_val.font.size = Pt(11)

    add_subtitulo("DATOS GENERALES", 1)
    for c in ["Código", "Fecha", "Hora", "Nombres y Apellidos", "DNI", "Edad", "Sexo", "Cargo", "Área", "Tiempo de servicio", "Tipo de contrato", "Teléfono"]:
        if c in datos: add_texto(c, datos[c])

    add_subtitulo("MOTIVO DE ATENCIÓN", 2)
    for c in ["Motivo de Atención", "Solicitante"]:
        if c in datos: add_texto(c, datos[c])

    add_subtitulo("DESCRIPCIÓN DEL PROBLEMA", 3)
    for c in ["Descripción", "Tiempo del problema", "Ámbito del problema"]:
        if c in datos: add_texto(c, datos[c])

    add_subtitulo("OBSERVACIÓN DEL PSICÓLOGO", 4)
    for c in ["Actitud", "Observaciones conductuales", "Área afectada"]:
        if c in datos: add_texto(c, datos[c])

    add_subtitulo("ORIENTACIÓN Y CONSEJERÍA", 5)
    for c in ["Orientación", "Acuerdos"]:
        if c in datos: add_texto(c, datos[c])

    add_subtitulo("PLAN DE ACCIÓN", 6)
    for c in ["Plan de Acción", "Requiere cita", "Fecha próxima cita"]:
        if c in datos: add_texto(c, datos[c])

    add_subtitulo("CONCLUSIÓN Y RECOMENDACIONES", 7)
    for c in ["Conclusión"]:
        if c in datos: add_texto(c, datos[c])

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# -------------------------
# LOGIN
# -------------------------
def login():
    st.title("🔐 Acceso al Sistema")
    user = st.text_input("Usuario")
    password = st.text_input("Contraseña", type="password")

    if st.button("Ingresar"):
        if user == "psicologoluigi" and password == "psicologoluigi151297":
            st.session_state["login"] = True
            st.rerun()
        else:
            st.error("Credenciales incorrectas")

if "login" not in st.session_state:
    st.session_state["login"] = False

if not st.session_state["login"]:
    login()
    st.stop()

# -------------------------
# MENÚ
# -------------------------
st.sidebar.title("🧠 Menú")
menu = st.sidebar.radio("Navegación", 
                        ["🏠 Inicio", "📋 Nueva Atención", "📈 Seguimiento", "📂 Historial", "🔎 Buscar por DNI", "📊 Reportes y Alertas", "🚪 Cerrar sesión"])

df_atenciones = cargar_datos()

# -------------------------
# INICIO
# -------------------------
if menu == "🏠 Inicio":
    st.title("Sistema de Psicología Organizacional")
    st.write("Subgerencia de Seguridad Ciudadana")

    col1, col2, col3 = st.columns(3)
    col1.metric("Total Atenciones", len(df_atenciones))
    
    if not df_atenciones.empty:
        col2.metric("Trabajadores únicos", df_atenciones["DNI"].nunique())
        col3.metric("Última atención", df_atenciones["Fecha"].iloc[-1])
    else:
        col2.metric("Trabajadores únicos", 0)
        col3.metric("Última atención", "-")

# -------------------------
# NUEVA ATENCIÓN
# -------------------------
if menu == "📋 Nueva Atención":
    st.title("📋 Ficha de Atención Psicológica Laboral")

    fecha = datetime.now(ZONA_PERU).strftime("%Y-%m-%d")
    hora = datetime.now(ZONA_PERU).strftime("%H:%M:%S")

    dni_input = st.text_input("Ingrese DNI del trabajador para iniciar:")
    
    if dni_input:
        codigo, existe = obtener_codigo_por_dni(dni_input, df_atenciones)
        if existe:
            st.success(f"📁 Trabajador con historial previo - Código HC recuperado: {codigo}")
        else:
            st.info(f"🆕 Nuevo registro - Código HC asignado: {codigo}")

        with st.form("ficha", clear_on_submit=True):
            st.subheader("1. Datos Generales")
            
            st.text_input("Código", value=codigo, disabled=True)
            dni_form = st.text_input("DNI (Confirmar)", value=dni_input)

            col1, col2 = st.columns(2)
            with col1:
                nombre = st.text_input("Nombres y Apellidos")
                edad = st.number_input("Edad", 18, 70)
                sexo = st.selectbox("Sexo", ["Masculino","Femenino","Otro"])
                cargo = st.text_input("Cargo")
                area = st.selectbox("Área / Base", ["Serenazgo / Patrullaje", "Centro de Monitoreo / CCTV", "Guardianía / Puestos Fijos", "Administrativo", "Otro"])
            with col2:
                tiempo = st.text_input("Tiempo de servicio")
                contrato = st.selectbox("Tipo de contrato", ["CAS", "Permanente", "Nombrado", "Locador", "Otro"])
                telefono = st.text_input("Teléfono")

            st.subheader("2. Motivo de Atención")
            
            opciones_oficiales = [
                "Estrés laboral", "Conflicto con compañero", "Conflicto con jefe", "Problemas familiares", 
                "Problemas de pareja", "Problemas económicos", "Desmotivación laboral", "Problemas de conducta laboral", 
                "Dificultad de adaptación al trabajo", "Problemas de trabajo en equipo", "Problemas disciplinarios", 
                "Bajo rendimiento laboral", "Problemas personales", "Orientación laboral", "Otros"
            ]
            
            if area == "Serenazgo / Patrullaje":
                motivos_area = ["Estrés laboral (Agresiones en calle)", "Falta de respaldo en intervenciones"] + opciones_oficiales
            elif area == "Centro de Monitoreo / CCTV":
                motivos_area = ["Trauma vicario (Visualización de accidentes/delitos)", "Fatiga visual / mental extrema"] + opciones_oficiales
            elif area == "Guardianía / Puestos Fijos":
                motivos_area = ["Aislamiento / Soledad", "Monotonía extrema", "Problemas ergonómicos"] + opciones_oficiales
            elif area == "Administrativo":
                motivos_area = ["Sobrecarga de atención al público", "Estrés burocrático"] + opciones_oficiales
            else:
                motivos_area = opciones_oficiales
            
            motivos_area = list(dict.fromkeys(motivos_area))
            motivo = st.selectbox("Motivo principal", motivos_area)
            
            motivo_otro = ""
            if motivo == "Otros" or "Otros:" in motivo:
                motivo_otro = st.text_input("Especifique otro motivo:")

            solicitante = st.selectbox("¿Quién solicita la atención?", ["Voluntario", "Derivado por jefe", "Recursos Humanos", "Psicología (seguimiento)", "Otro"])

            st.subheader("3. Descripción del Problema")
            descripcion = st.text_area("Descripción del problema")
            
            col_t1, col_t2 = st.columns(2)
            with col_t1:
                tiempo_prob = st.selectbox("Tiempo del problema", ["Días", "Semanas", "Meses", "Años"])
            with col_t2:
                ambito = st.selectbox("Ámbito del problema", ["Laboral", "Familiar", "Personal", "Económico", "Pareja", "Mixto"])

            st.subheader("4. Observación del Psicólogo")
            actitud = st.selectbox("Actitud durante la entrevista", ["Colaborador", "Reservado", "Evasivo", "Agresivo", "Ansioso", "Desmotivado", "Preocupado", "Otro"])
            observaciones = st.text_area("Observaciones conductuales")
            area_afectada = st.multiselect("Área afectada", ["Desempeño laboral", "Relaciones laborales", "Puntualidad", "Trabajo en equipo", "Trato al ciudadano", "Cumplimiento de órdenes", "Ninguna", "Otros"])

            st.subheader("5. Orientación / Consejería Brindada")
            orientacion = st.text_area("Orientación / Consejería")
            acuerdos = st.text_area("Acuerdos o compromisos")

            st.subheader("6. Plan de Acción")
            plan = st.multiselect("Plan de Acción", ["Seguimiento psicológico laboral", "Derivación a Recursos Humanos", "Derivación a jefe inmediato", "Recomendación de capacitación", "Mediación laboral", "Sin seguimiento", "Otros"])
            
            # --- SOLUCIÓN A LA PRÓXIMA CITA EN NUEVA ATENCIÓN ---
            col_cita1, col_cita2 = st.columns(2)
            with col_cita1:
                requiere_cita = st.radio("¿Requiere próxima cita?", ["No", "Sí"])
            with col_cita2:
                fecha_prox = st.date_input("Seleccione la fecha (Solo si requiere)")

            st.subheader("7. Conclusión y Recomendaciones Laborales")
            conclusion = st.text_area("Conclusión y recomendaciones")

            guardar = st.form_submit_button("💾 Guardar Atención")

        if guardar:
            motivo_final = motivo if motivo != "Otros" else f"Otros: {motivo_otro}"
            
            # Aquí la magia: Si el doctor dijo "No", ignoramos la fecha ingresada
            fecha_prox_str = str(fecha_prox) if requiere_cita == "Sí" else "No requiere"
            
            dni_final = dni_form.strip().zfill(8)

            datos = {
                "Código": codigo, "Fecha": fecha, "Hora": hora, "Nombres y Apellidos": nombre, "DNI": dni_final,
                "Edad": edad, "Sexo": sexo, "Cargo": cargo, "Área": area, "Tiempo de servicio": tiempo,
                "Tipo de contrato": contrato, "Teléfono": telefono, "Motivo de Atención": motivo_final, "Solicitante": solicitante,
                "Descripción": descripcion, "Tiempo del problema": tiempo_prob, "Ámbito del problema": ambito,
                "Actitud": actitud, "Observaciones conductuales": observaciones, "Área afectada": ", ".join(area_afectada),
                "Orientación": orientacion, "Acuerdos": acuerdos, "Plan de Acción": ", ".join(plan),
                "Requiere cita": requiere_cita, "Fecha próxima cita": fecha_prox_str, "Conclusión": conclusion
            }

            if guardar_datos(datos):
                st.success(f"✅ Atención guardada exitosamente en la nube.")
            
            archivo_word_memoria = generar_word_memoria(datos)
            st.download_button(
                label="📄 Descargar ficha en Word",
                data=archivo_word_memoria,
                file_name=f"Ficha_{codigo}_{fecha}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# -------------------------
# SEGUIMIENTO
# -------------------------
if menu == "📈 Seguimiento":
    st.title("📈 Seguimiento de Caso")
    dni_seg = st.text_input("Ingrese DNI para seguimiento:")
    
    if dni_seg:
        if not df_atenciones.empty:
            dni_busqueda = dni_seg.strip().zfill(8)
            resultado = df_atenciones[df_atenciones["DNI"] == dni_busqueda]
            
            if not resultado.empty:
                ultimo_registro = resultado.iloc[-1]
                st.success(f"Paciente: {ultimo_registro['Nombres y Apellidos']} | Código: {ultimo_registro['Código']}")
                st.info(f"Último motivo: {ultimo_registro.get('Motivo de Atención', 'No registrado')}")

                with st.form("ficha_seguimiento", clear_on_submit=True):
                    st.subheader("Evolución del Caso")
                    col1, col2 = st.columns(2)
                    with col1:
                        estado_evolucion = st.selectbox("Estado actual", ["Mejoría notable", "Mejoría leve", "Estable/Sin cambios", "Retroceso/Empeoramiento"])
                        cumplimiento = st.select_slider("Cumplimiento de acuerdos previos", options=["Ninguno", "Bajo", "Medio", "Alto", "Total"])
                    with col2:
                        nueva_actitud = st.selectbox("Actitud en sesión", ["Colaborador", "Reservado", "Evasivo", "Agresivo", "Ansioso", "Desmotivado", "Preocupado", "Otro"])
                        
                        fecha_hoy = datetime.now(ZONA_PERU).strftime("%Y-%m-%d")
                        hora_hoy = datetime.now(ZONA_PERU).strftime("%H:%M:%S")

                    st.subheader("Detalles de la sesión")
                    evolucion_detallada = st.text_area("Descripción de la evolución (¿Qué ha pasado desde la última vez?)")
                    nuevos_acuerdos = st.text_area("Nuevos acuerdos o compromisos")
                    
                    st.subheader("Plan de Acción")
                    proximo_paso = st.multiselect("Plan de Acción (Seguimiento)", ["Seguimiento psicológico laboral", "Derivación a Recursos Humanos", "Alta administrativa", "Derivación externa", "Cierre de caso", "Otros"])
                    
                    # --- SOLUCIÓN A LA PRÓXIMA CITA EN SEGUIMIENTO ---
                    col_seg1, col_seg2 = st.columns(2)
                    with col_seg1:
                        req_cita_seg = st.radio("¿Nueva próxima cita?", ["No", "Sí"])
                    with col_seg2:
                        fecha_prox_seg = st.date_input("Fecha sugerida (Solo si requiere)")

                    guardar_seg = st.form_submit_button("💾 Registrar Seguimiento")

                if guardar_seg:
                    # Aplicando la misma magia
                    fecha_prox_seg_str = str(fecha_prox_seg) if req_cita_seg == "Sí" else "No requiere"
                    
                    datos_seg = {
                        "Código": ultimo_registro['Código'], "Fecha": fecha_hoy, "Hora": hora_hoy,
                        "Nombres y Apellidos": ultimo_registro['Nombres y Apellidos'], "DNI": dni_busqueda,
                        "Edad": ultimo_registro['Edad'], "Sexo": ultimo_registro['Sexo'], "Cargo": ultimo_registro['Cargo'],
                        "Área": ultimo_registro['Área'], "Tiempo de servicio": ultimo_registro.get('Tiempo de servicio', '-'),
                        "Tipo de contrato": ultimo_registro.get('Tipo de contrato', '-'), "Teléfono": ultimo_registro.get('Teléfono', '-'),
                        "Motivo de Atención": f"SEGUIMIENTO: {ultimo_registro.get('Motivo de Atención', '')}",
                        "Solicitante": "Psicología (seguimiento)", "Descripción": evolucion_detallada,
                        "Tiempo del problema": "-", "Ámbito del problema": "-", "Actitud": nueva_actitud,
                        "Observaciones conductuales": "-", "Área afectada": "-", "Orientación": "-", "Acuerdos": nuevos_acuerdos,
                        "Plan de Acción": ", ".join(proximo_paso), "Requiere cita": req_cita_seg,
                        "Fecha próxima cita": fecha_prox_seg_str, "Conclusión": f"Estado: {estado_evolucion}. Cumplimiento acuerdos: {cumplimiento}"
                    }
                    
                    if guardar_datos(datos_seg):
                        st.success("✅ Seguimiento registrado con éxito.")
                        
                    doc_seg = generar_word_memoria(datos_seg)
                    st.download_button("📄 Descargar Ficha de Seguimiento", doc_seg, file_name=f"Seguimiento_{dni_busqueda}_{fecha_hoy}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                st.error("No se encontró ningún registro previo con ese DNI.")
        else:
            st.error("La base de datos está vacía.")

# -------------------------
# HISTORIAL Y BÚSQUEDA
# -------------------------
if menu == "📂 Historial":
    st.title("Historial General")
    st.dataframe(df_atenciones, use_container_width=True)

if menu == "🔎 Buscar por DNI":
    st.title("Buscar paciente")
    dni_buscar = st.text_input("Ingrese DNI")
    if st.button("Buscar"):
        if not df_atenciones.empty:
            dni_busqueda = dni_buscar.strip().zfill(8)
            resultado = df_atenciones[df_atenciones["DNI"] == dni_busqueda]
            
            if not resultado.empty:
                st.dataframe(resultado, use_container_width=True)
                
                st.markdown("---")
                st.subheader("📄 Descargar Fichas Clínicas")
                st.write("Se encontraron las siguientes atenciones. Seleccione la que desea descargar:")
                
                for index, row in resultado.iterrows():
                    datos_fila = row.to_dict()
                    archivo_word = generar_word_memoria(datos_fila)
                    
                    tipo_atencion = "Atención Primaria"
                    if "SEGUIMIENTO" in str(row.get('Motivo de Atención', '')):
                        tipo_atencion = "Seguimiento"
                        
                    st.download_button(
                        label=f"⬇️ Descargar Ficha de {tipo_atencion} del {row['Fecha']} ({row['Hora']})",
                        data=archivo_word,
                        file_name=f"Ficha_{row['Código']}_{row['Fecha']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"dl_btn_{index}" 
                    )
            else:
                st.warning("No se encontraron registros para ese DNI.")
        else:
            st.info("La base de datos está vacía.")

# -------------------------
# REPORTES Y ALERTAS (DASHBOARD)
# -------------------------
if menu == "📊 Reportes y Alertas":
    st.title("📊 Panel de Control Psicosocial")
    st.write("Análisis epidemiológico y alertas organizacionales.")

    if df_atenciones.empty:
        st.info("No hay datos suficientes para generar reportes.")
    else:
        df_base = df_atenciones[~df_atenciones["Motivo de Atención"].astype(str).str.contains("SEGUIMIENTO", na=False)]
        
        st.subheader("🚨 Alertas de Riesgo")
        col_alerta1, col_alerta2 = st.columns(2)
        
        with col_alerta1:
            conteo_pacientes = df_atenciones["DNI"].value_counts()
            pacientes_frecuentes = conteo_pacientes[conteo_pacientes >= 3]
            
            if not pacientes_frecuentes.empty:
                st.error(f"⚠️ **Atención Crítica:** Hay {len(pacientes_frecuentes)} paciente(s) con 3 o más atenciones. Requieren evaluación detallada.")
                for dni, atenciones in pacientes_frecuentes.items():
                    nombre_paciente = df_atenciones[df_atenciones["DNI"] == dni]["Nombres y Apellidos"].iloc[0]
                    st.write(f"- **{nombre_paciente}** (DNI: {dni}): {atenciones} visitas")
            else:
                st.success("✅ No hay pacientes con hiper-recurrencia (3+ visitas).")

        with col_alerta2:
            if not df_base.empty:
                area_critica = df_base["Área"].value_counts().idxmax()
                cantidad_area = df_base["Área"].value_counts().max()
                st.warning(f"🔥 **Área Caliente:** **{area_critica}** es el sector con más intervenciones ({cantidad_area} casos).")
            
        st.markdown("---")
        st.subheader("📈 Distribución de Casos")
        
        if not df_base.empty:
            col_graf1, col_graf2 = st.columns(2)

            with col_graf1:
                st.markdown("**Principales Motivos de Consulta**")
                conteo_motivos = df_base["Motivo de Atención"].value_counts().reset_index()
                conteo_motivos.columns = ['Motivo', 'Cantidad']
                
                fig_motivos = px.pie(
                    conteo_motivos, 
                    values='Cantidad', 
                    names='Motivo', 
                    hole=0.4, 
                    color_discrete_sequence=px.colors.sequential.RdBu
                )
                fig_motivos.update_traces(textposition='inside', textinfo='percent+label')
                st.plotly_chart(fig_motivos, use_container_width=True)

            with col_graf2:
                st.markdown("**Intervenciones por Área de Trabajo**")
                conteo_areas = df_base["Área"].value_counts().reset_index()
                conteo_areas.columns = ['Área', 'Cantidad']
                
                fig_areas = px.bar(
                    conteo_areas, 
                    x='Área', 
                    y='Cantidad', 
                    text='Cantidad',
                    color='Cantidad',
                    color_continuous_scale='Reds'
                )
                fig_areas.update_layout(xaxis_title="Área / Base", yaxis_title="Número de Casos")
                st.plotly_chart(fig_areas, use_container_width=True)
                
            st.subheader("📋 Detalle de Motivos por Área")
            tabla_cruzada = pd.crosstab(df_base['Área'], df_base['Motivo de Atención'])
            st.dataframe(tabla_cruzada, use_container_width=True)
            
        else:
            st.info("Registra nuevas atenciones (no seguimientos) para visualizar los gráficos.")

# -------------------------
# LOGOUT
# -------------------------
if menu == "🚪 Cerrar sesión":
    st.session_state["login"] = False
    st.rerun()
